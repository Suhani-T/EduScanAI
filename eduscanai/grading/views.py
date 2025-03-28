from django.shortcuts import render
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.core.mail import EmailMessage, send_mail
import google.generativeai as genai
import environ
import json
import chardet
import easyocr
from pdf2image import convert_from_bytes
import tempfile
import fitz
import pytesseract
from docx import Document
import markdown
from google.cloud import documentai_v1beta3 as documentai
from google.cloud import storage
import time



import random
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from django.core.cache import cache 


from django.conf import settings
import os

env = environ.Env()
environ.Env.read_env()

API_KEY = env("GEMINI_API_KEY", default=None)

if API_KEY:
    genai.configure(api_key=API_KEY)
else:
    print(" gemini API error")

def index(request):
    return render(request, "index.html")  


def extract_text_from_pdf(pdf_bytes):
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    extracted_text = "\n".join(page.get_text("text") for page in doc)
    return extracted_text

def extract_text_from_docx(docx_bytes):
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp.write(docx_bytes)
        tmp_path = tmp.name

    doc = Document(tmp_path)
    text = "\n".join([para.text for para in doc.paragraphs])

    os.remove(tmp_path)  # cleanup the temp file
    return text

    # with open("temp.docx", "wb") as f:
    #     f.write(docx_bytes)

    # doc = Document("temp.docx")
    # return "\n".join([para.text for para in doc.paragraphs])

def extract_text(file):
    file_name = file.name.lower()
    file_bytes = file.read()

    if file_name.endswith(".txt"):
        encoding = chardet.detect(file_bytes)["encoding"]
        return file_bytes.decode(encoding or "utf-8", errors="ignore")

    elif file_name.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)

    elif file_name.endswith(".docx"):
        return extract_text_from_docx(file_bytes)

    else:
        return None  



@csrf_exempt
def evaluate(request):
    if request.method == "POST":
        try:
            answer_key_file = request.FILES.get("answer_key")
            student_script_file = request.FILES.get("student_script")

            if not answer_key_file or not student_script_file:
                return JsonResponse({"error": "Both answer key and student script are required."}, status=400)

            answer_key = extract_text(answer_key_file)
            student_script = extract_text(student_script_file)

            if not answer_key or not student_script:
                return JsonResponse({"error": "Failed to extract text from one or both files."}, status=400)

            print("extracted answer key:\n", answer_key[:500])  
            print("extracted student script:\n", student_script[:500])

            model = genai.GenerativeModel("gemini-1.5-flash")
            response = model.generate_content(f"Compare this student answer: {student_script} with answer key: {answer_key}. Provide constructive feedback.")

            feedback = response.candidates[0].content.parts[0].text if response.candidates else "Unable to generate feedback."

            return JsonResponse({"feedback": feedback})

        except Exception as e:
            print("error in evaluate():", str(e))
            return JsonResponse({"error": "Failed to generate feedback", "details": str(e)}, status=500)

@csrf_exempt
def evaluate(request):
    if request.method == "POST":
        try:
            answer_key_file = request.FILES.get("answer_key")
            student_script_file = request.FILES.get("student_script")

            custom_prompt = request.POST.get("custom_prompt", "").strip()


            if not answer_key_file or not student_script_file:
                return JsonResponse({"error": "Both answer key and student script are required."}, status=400)

            answer_key = extract_text(answer_key_file)
            student_script = extract_text(student_script_file)

            if not answer_key or not student_script:
                return JsonResponse({"error": "Failed to extract text from one or both files."}, status=400)

            default_prompt = f"""
                Evaluate the student's answer: {student_script} against the answer key: {answer_key} based on the following criteria:
                1. Constructive Feedback: Highlight correct points, errors, and areas for improvement.
                2. Score Calculation: Assign a total score based on the marking scheme, if available. Apply partial marking for subjective questions unless explicitly stated by the teacher.
                3. Spelling Considerations: Ignore minor spelling mistakes, except in language, grammar, or literature subjects.
                4. Answer Relevance Check: Determine if the response fully, partially, or does not address the question. Highlight missing key points.
                5. Clarity & Coherence Assessment: Evaluate whether the answer is well-structured and logically presented.
                6. Keyword Matching (for factual/technical subjects): Ensure key terms or concepts from the answer key are present.
                7. Difficulty Level Analysis: Identify if the student struggles with basic or advanced concepts based on errors.
                8. Misconception Detection: Spot and correct misunderstandings or misinterpretations.
                9. Response Comparison: Compare the student’s response against high-quality sample answers.
                10. Study Resources: Suggest relevant materials for improvement.
                11. Format Output as follows: 
                    - Feedback: Strengths, mistakes, and improvement suggestions.  
                    - Score: Marks obtained for that specific question.  
                    - Study Resources: Recommendations for weak areas.
                12. Provide a brief summary of the student's overall strengths and weaknesses.
                13. End the feedback with 2-3 motivational or encouraging sentences naturally, without using headings like 'Conclusion' or 'Concluding Note'. It should feel like a genuine and warm closing remark.
            """
            if custom_prompt:
                prompt = f"Teacher's Custom Evaluation Criteria (takes precedence over the default rules if conflicting): \n {custom_prompt} \n\n Default Rules: \n{default_prompt}"
            else:
                prompt = default_prompt

            # final_prompt = f"{custom_prompt}\n\n{default_prompt}" if custom_prompt else default_prompt

            model = genai.GenerativeModel("gemini-1.5-flash")
            response = model.generate_content(prompt)

            feedback = response.candidates[0].content.parts[0].text if response.candidates else "Unable to generate feedback."

            return JsonResponse({"feedback": feedback})

        except Exception as e:
            print("Error in evaluate():", str(e))
            return JsonResponse({"error": "Failed to generate feedback", "details": str(e)}, status=500)

@csrf_exempt
def send_feedback(request):
    if request.method == "POST":
        try:
            data = json.loads(request.body)
            student_email = data.get("student_email")
            feedback_markdown = data.get("feedback")  
            teacher_email = data.get("teacher_email")

            
            feedback_html = markdown.markdown(feedback_markdown)

            email = EmailMessage(
                subject="Feedback on Your Answer Script",
                body=feedback_html,
                from_email="aieduscan@gmail.com",
                to=[student_email],
                headers={"Reply-To": teacher_email}
            )
            email.content_subtype = "html"  
            email.send()

            return JsonResponse({"message": "feedback sent successfully"})
        
        except Exception as e:
            print("error in send_feedback():", str(e))
            return JsonResponse({"error": "failed to send feedback", "details": str(e)}, status=500)
        

def is_scanned_pdf(pdf_bytes):
   
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    for page in doc:
        if page.get_text(): 
            return False
    return True       

PROJECT_ID = "eduscanai"
LOCATION = "us"
PROCESSOR_ID = "6db65a3b9bc94c91"  

def extract_text_with_document_ai(pdf_bytes):

    client = documentai.DocumentUnderstandingServiceClient()

    
    gcs_client = storage.Client()
    bucket_name = "your-gcs-bucket"
    bucket = gcs_client.bucket(bucket_name)
    blob = bucket.blob(f"scanned_pdfs/{time.time()}.pdf")
    blob.upload_from_string(pdf_bytes, content_type="application/pdf")
    
    
    gcs_uri = f"gs://{bucket_name}/{blob.name}"

    
    request = {
        "name": f"projects/{PROJECT_ID}/locations/{LOCATION}/processors/{PROCESSOR_ID}",
        "raw_document": {
            "content": pdf_bytes,
            "mime_type": "application/pdf",
        },
    }

    
    response = client.process_document(request=request)
    
    
    extracted_text = []
    for page in response.document.pages:
        for block in page.blocks:
            for paragraph in block.paragraphs:
                for word in paragraph.words:
                    extracted_text.append("".join([symbol.text for symbol in word.symbols]))

    return " ".join(extracted_text)



def extract_text(file):
    file_name = file.name.lower()
    file_bytes = file.read()

    if file_name.endswith(".txt"):
        encoding = chardet.detect(file_bytes)["encoding"]
        return file_bytes.decode(encoding or "utf-8", errors="ignore")

    elif file_name.endswith(".pdf"):
        if is_scanned_pdf(file_bytes):  
            return extract_text_with_document_ai(file_bytes)
        else:
            return extract_text_from_pdf(file_bytes)  

    elif file_name.endswith(".docx"):
        return extract_text_from_docx(file_bytes)

    else:
        return None  
    





def send_email(to_email, otp):
   
    from_email = os.environ.get('EMAIL_HOST_USER')  
    password = os.environ.get('EMAIL_HOST_PASSWORD') 
    
    try:
        server = smtplib.SMTP("smtp.gmail.com", 587)
        server.starttls()
        server.login(from_email, password)

        msg = MIMEMultipart()
        msg["From"] = from_email
        msg["To"] = to_email
        msg["Subject"] = "Your OTP for Authentication"

        body = f"Your OTP is {otp}. Please enter it in the website to verify your identity."
        msg.attach(MIMEText(body, "plain"))

        server.sendmail(from_email, to_email, msg.as_string())
        server.quit()
    except Exception as e:
        print(f"Failed to send email: {e}")




def send_otp(request):
    if request.method == "POST":
        try:
            
            email = request.POST.get("email")
            if not email:
                return JsonResponse({"status": "fail", "message": "Email is required."})

           
            otp = random.randint(100000, 999999)
            cache.set(email, otp, timeout=300)  

           
            send_email(email, otp)

            return JsonResponse({"status": "success", "message": "OTP sent successfully."})
        except Exception as e:
            return JsonResponse({"status": "fail", "message": f"Error: {e}"})

    return JsonResponse({"status": "fail", "message": "Invalid request method."})

@csrf_exempt
def verify_otp(request):
    if request.method == "POST":
        email = request.POST.get("email")
        otp = request.POST.get("otp")

       
        stored_otp = cache.get(email)

        if stored_otp and int(otp) == stored_otp:
            cache.delete(email)  
            return JsonResponse({"status": "success", "message": "OTP verified successfully."})
        else:
            return JsonResponse({"status": "fail", "message": "Invalid OTP."})

    return JsonResponse({"status": "fail", "message": "Invalid request method."})




from google.cloud import translate_v2 as translate

def translate_text(request):
    if request.method == 'POST':
        import json
        data = json.loads(request.body)
        text_list = data.get('text', [])
        target_lang = data.get('target', 'en')

        if not text_list:
            return JsonResponse({"error": "No text provided"}, status=400)

        translate_client = translate.Client()
        translations = translate_client.translate(text_list, target_language=target_lang)

        translated_texts = [t['translatedText'] for t in translations]

        return JsonResponse({"translations": translated_texts})
